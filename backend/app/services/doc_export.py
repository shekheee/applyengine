from __future__ import annotations

import io

from docx import Document
from docx.shared import Pt


def text_to_docx(text: str, title: str = "") -> bytes:
    """Render plain text into a simple, ATS-friendly .docx document."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        heading = doc.add_paragraph()
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

    for block in text.split("\n"):
        doc.add_paragraph(block)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
