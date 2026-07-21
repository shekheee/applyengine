from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models import Memory, Profile
from app.services.resume_designed import prepare_designed_resume_document
from app.services.resume_pdf import _safe_filename

_ACCENT = RGBColor(0x43, 0x38, 0xCA)
_HEAD = RGBColor(0x1A, 0x1A, 0x2E)
_MUTED = RGBColor(0x4B, 0x55, 0x63)


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _ACCENT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def render_resume_docx(doc: dict[str, Any]) -> bytes:
    """Build a well-structured .docx for Google Docs import."""
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    name = str(doc.get("name") or "Candidate")
    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = _HEAD
    name_p.paragraph_format.space_after = Pt(2)

    contact_parts = [
        p for p in [doc.get("email"), doc.get("phone"), doc.get("location")] if p
    ]
    links = doc.get("links") or []
    contact_parts.extend(str(link) for link in links[:2])
    if contact_parts:
        cp = document.add_paragraph(" · ".join(str(x) for x in contact_parts))
        cp.paragraph_format.space_after = Pt(6)
        for run in cp.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = _MUTED

    summary = (doc.get("summary") or "").strip()
    if summary:
        _add_section_heading(document, "Summary")
        sp = document.add_paragraph(summary)
        sp.paragraph_format.space_after = Pt(4)

    skills = doc.get("skills") or []
    if skills:
        _add_section_heading(document, "Skills")
        sk = document.add_paragraph(", ".join(str(s) for s in skills))
        sk.paragraph_format.space_after = Pt(4)

    experience = doc.get("experience") or []
    if experience:
        _add_section_heading(document, "Experience")
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            title = str(exp.get("title") or "").strip()
            company = str(exp.get("company") or "").strip()
            dates = str(exp.get("dates") or "").strip()
            head = " — ".join(x for x in [title, company] if x)
            if head:
                rp = document.add_paragraph()
                hr = rp.add_run(head)
                hr.bold = True
                hr.font.size = Pt(10.5)
                if dates:
                    dr = rp.add_run(f"  ({dates})")
                    dr.italic = True
                    dr.font.size = Pt(9)
                    dr.font.color.rgb = _MUTED
                rp.paragraph_format.space_before = Pt(4)
                rp.paragraph_format.space_after = Pt(1)
            for bullet in exp.get("highlights") or []:
                bp = document.add_paragraph(str(bullet), style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
                for run in bp.runs:
                    run.font.size = Pt(10)

    projects = doc.get("projects") or []
    if projects:
        _add_section_heading(document, "Projects")
        for proj in projects:
            if not isinstance(proj, dict):
                continue
            pname = str(proj.get("name") or "Project")
            desc = str(proj.get("description") or "")
            tech = proj.get("tech") or []
            pp = document.add_paragraph()
            pr = pp.add_run(pname)
            pr.bold = True
            if desc:
                pp.add_run(f" — {desc}")
            if tech:
                tp = document.add_paragraph(
                    "Tools & methods: " + ", ".join(str(t) for t in tech[:8])
                )
                for run in tp.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = _MUTED

    education = doc.get("education") or []
    if education:
        _add_section_heading(document, "Education")
        for ed in education:
            if not isinstance(ed, dict):
                continue
            school = str(ed.get("school") or "").strip()
            degree = str(ed.get("degree") or "").strip()
            dates = str(ed.get("dates") or "").strip()
            line = " — ".join(x for x in [degree, school] if x) or school
            ep = document.add_paragraph()
            er = ep.add_run(line)
            er.bold = True
            if dates:
                ep.add_run(f" ({dates})")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_resume_docx_from_doc(doc: dict[str, Any]) -> tuple[bytes, str]:
    docx_bytes = render_resume_docx(doc)
    fname = f"{_safe_filename(str(doc.get('name', 'resume')))}_resume.docx"
    return docx_bytes, fname


def build_resume_docx(
    profile: Profile | None,
    memories: list[Memory],
    job=None,
) -> tuple[bytes, str]:
    doc = prepare_designed_resume_document(profile, memories, job)
    docx_bytes = render_resume_docx(doc)
    fname = f"{_safe_filename(str(doc.get('name', 'resume')))}_resume.docx"
    return docx_bytes, fname
